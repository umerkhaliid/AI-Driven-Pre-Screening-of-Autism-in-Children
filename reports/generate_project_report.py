from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
IMG_DIR = OUT_DIR / "chapter3_images"
DOCX_PATH = OUT_DIR / "Autism_PreScreening_Project_Report.docx"


def _setup_canvas(width: float = 14, height: float = 8):
    fig, ax = plt.subplots(figsize=(width, height))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")
    return fig, ax


def _box(ax, x, y, w, h, text, fc="#EAF2FF", ec="#2F4F8F"):
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=1.5",
        linewidth=1.5,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=10)


def _arrow(ax, x1, y1, x2, y2, text=""):
    ax.annotate(
        "",
        xy=(x2, y2),
        xytext=(x1, y1),
        arrowprops=dict(arrowstyle="->", lw=1.4, color="#1D3557"),
    )
    if text:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 1.5, text, fontsize=9, ha="center")


def create_architecture():
    fig, ax = _setup_canvas()
    ax.set_title("Figure 1: System Architecture of Autism Pre-Screening Tool", fontsize=15, weight="bold")
    _box(ax, 3, 42, 14, 16, "Parent /\nCaregiver")
    _box(ax, 23, 60, 20, 14, "Web Frontend\n(web/public)")
    _box(ax, 23, 28, 20, 14, "Streamlit UI\n(app/ui)")
    _box(ax, 50, 44, 20, 16, "FastAPI Backend\n(server/app.py)")
    _box(ax, 76, 70, 20, 12, "Face Screening\n(src/face_screening.py)")
    _box(ax, 76, 54, 20, 12, "Screening Inference\n(src/inference.py)")
    _box(ax, 76, 38, 20, 12, "LLM Report\n(src/llm_report_groq.py)")
    _box(ax, 76, 22, 20, 12, "PDF Generator\n(src/pdf_generator.py)")
    _box(ax, 50, 10, 20, 12, "Model Artifacts\n(models/*.joblib, *.pth)")
    _box(ax, 76, 6, 20, 12, "External LLM API\n(Groq)")

    _arrow(ax, 17, 50, 23, 67, "Uses")
    _arrow(ax, 17, 50, 23, 35, "Optional")
    _arrow(ax, 43, 67, 50, 52, "HTTP")
    _arrow(ax, 43, 35, 50, 52, "HTTP")
    _arrow(ax, 70, 52, 76, 76)
    _arrow(ax, 70, 52, 76, 60)
    _arrow(ax, 70, 52, 76, 44)
    _arrow(ax, 70, 52, 76, 28)
    _arrow(ax, 70, 48, 50, 16, "Load/Infer")
    _arrow(ax, 86, 44, 86, 18, "Report -> PDF")
    _arrow(ax, 86, 38, 86, 12, "LLM call")

    fig.tight_layout()
    fig.savefig(IMG_DIR / "chapter3_architecture.png", dpi=220)
    plt.close(fig)


def create_usecase():
    fig, ax = _setup_canvas()
    ax.set_title("Figure 2: Use Case Diagram (Functional View)", fontsize=15, weight="bold")
    _box(ax, 3, 70, 15, 10, "Parent/\nCaregiver", fc="#FFF3E6", ec="#A35D00")
    _box(ax, 3, 20, 15, 10, "System\nAdmin", fc="#FFF3E6", ec="#A35D00")
    _box(ax, 25, 8, 70, 84, "Autism Pre-Screening System", fc="#F8FBFF", ec="#2F4F8F")

    use_cases = [
        (32, 72, "Capture Child Photo"),
        (57, 72, "Complete Questionnaire"),
        (32, 56, "Run Face Screening"),
        (57, 56, "Generate Risk Prediction"),
        (32, 40, "Generate LLM Report"),
        (57, 40, "Download PDF Report"),
        (44, 24, "Check API Health / Logs"),
    ]
    for x, y, label in use_cases:
        _box(ax, x, y, 28, 10, label, fc="#E8F6EF", ec="#1E7D4E")

    _arrow(ax, 18, 75, 32, 77)
    _arrow(ax, 18, 75, 57, 77)
    _arrow(ax, 18, 75, 32, 61)
    _arrow(ax, 18, 75, 57, 61)
    _arrow(ax, 18, 75, 32, 45)
    _arrow(ax, 18, 75, 57, 45)
    _arrow(ax, 18, 25, 44, 29)

    fig.tight_layout()
    fig.savefig(IMG_DIR / "chapter3_usecase.png", dpi=220)
    plt.close(fig)


def create_sequence():
    fig, ax = _setup_canvas(height=9)
    ax.set_title("Figure 3: Screening and Reporting Sequence", fontsize=15, weight="bold")
    lifelines = [
        (10, "Parent"),
        (28, "Frontend"),
        (46, "FastAPI"),
        (64, "Face Service"),
        (80, "Inference Service"),
        (92, "LLM/PDF"),
    ]
    for x, name in lifelines:
        ax.text(x, 95, name, ha="center", fontsize=10, weight="bold")
        ax.plot([x, x], [15, 92], linestyle="--", linewidth=1, color="gray")

    y = 86
    steps = [
        (10, 28, "Upload photo"),
        (28, 46, "POST /api/photo/screen"),
        (46, 64, "predict_face_binary()"),
        (64, 46, "face result"),
        (46, 28, "autistic? continue"),
        (10, 28, "submit questionnaire"),
        (28, 46, "POST /api/screen/predict"),
        (46, 80, "predict_autism_risk()"),
        (80, 46, "risk + probabilities"),
        (28, 46, "POST /api/report/llm"),
        (46, 92, "generate_risk_report()"),
        (92, 46, "report text"),
        (28, 46, "POST /api/report/pdf"),
        (46, 92, "generate_pdf_report()"),
        (92, 28, "PDF file response"),
    ]
    for src, dst, txt in steps:
        _arrow(ax, src, y, dst, y, txt)
        y -= 4.5

    fig.tight_layout()
    fig.savefig(IMG_DIR / "chapter3_sequence.png", dpi=220)
    plt.close(fig)


def create_activity():
    fig, ax = _setup_canvas()
    ax.set_title("Figure 4: Activity Diagram of User Workflow", fontsize=15, weight="bold")
    _box(ax, 8, 82, 16, 8, "Start", fc="#E6FFF2", ec="#2E8B57")
    _box(ax, 32, 82, 24, 8, "Open UI and Enter Child Info")
    _box(ax, 62, 82, 26, 8, "Capture/Upload Photo")
    _box(ax, 62, 66, 26, 8, "Run Face Screening")
    _box(ax, 62, 50, 26, 8, "Face Result Decision")
    _box(ax, 8, 50, 40, 8, "If non-autistic: show guidance and stop")
    _box(ax, 62, 34, 26, 8, "If autistic: answer Q-CHAT/M-CHAT")
    _box(ax, 62, 18, 26, 8, "Generate Risk Prediction")
    _box(ax, 32, 18, 24, 8, "Generate LLM & PDF Report")
    _box(ax, 8, 18, 16, 8, "End", fc="#E6FFF2", ec="#2E8B57")

    _arrow(ax, 24, 86, 32, 86)
    _arrow(ax, 56, 86, 62, 86)
    _arrow(ax, 75, 82, 75, 74)
    _arrow(ax, 75, 66, 75, 58)
    _arrow(ax, 62, 54, 48, 54, "No")
    _arrow(ax, 75, 50, 75, 42, "Yes")
    _arrow(ax, 75, 34, 75, 26)
    _arrow(ax, 62, 22, 56, 22)
    _arrow(ax, 32, 22, 24, 22)

    fig.tight_layout()
    fig.savefig(IMG_DIR / "chapter3_activity.png", dpi=220)
    plt.close(fig)


def create_component():
    fig, ax = _setup_canvas()
    ax.set_title("Figure 5: Component Diagram", fontsize=15, weight="bold")
    _box(ax, 7, 72, 24, 12, "Presentation Layer\nindex.html / app.mjs / Streamlit")
    _box(ax, 38, 72, 24, 12, "API Layer\nFastAPI Routes")
    _box(ax, 69, 72, 24, 12, "Business Logic\ninference, scoring, mapping")
    _box(ax, 7, 46, 24, 12, "Face Module\nCNN calibrated pipeline")
    _box(ax, 38, 46, 24, 12, "Report Module\nLLM + PDF generation")
    _box(ax, 69, 46, 24, 12, "Model/Data Layer\njoblib, pth, datasets")
    _box(ax, 38, 20, 24, 12, "External Services\nGroq API")

    _arrow(ax, 31, 78, 38, 78)
    _arrow(ax, 62, 78, 69, 78)
    _arrow(ax, 50, 72, 50, 58)
    _arrow(ax, 81, 72, 81, 58)
    _arrow(ax, 50, 46, 50, 32)
    _arrow(ax, 19, 72, 19, 58)
    _arrow(ax, 31, 52, 38, 52)
    _arrow(ax, 62, 52, 69, 52)

    fig.tight_layout()
    fig.savefig(IMG_DIR / "chapter3_component.png", dpi=220)
    plt.close(fig)


def create_deployment():
    fig, ax = _setup_canvas()
    ax.set_title("Figure 6: Deployment Diagram", fontsize=15, weight="bold")
    _box(ax, 8, 64, 24, 20, "Client Device\nBrowser / Mobile")
    _box(ax, 40, 64, 24, 20, "Application Server\nPython + FastAPI\nStatic frontend")
    _box(ax, 72, 64, 20, 20, "Model Storage\nmodels/, data/")
    _box(ax, 40, 30, 24, 20, "Inference Runtime\nscikit-learn, torch")
    _box(ax, 72, 30, 20, 20, "Groq Cloud API\nLLM report")
    _box(ax, 8, 30, 24, 20, "Generated Outputs\nJSON results, PDF reports")

    _arrow(ax, 32, 74, 40, 74, "HTTPS")
    _arrow(ax, 64, 74, 72, 74, "Load artifacts")
    _arrow(ax, 52, 64, 52, 50, "invoke")
    _arrow(ax, 64, 40, 72, 40, "API call")
    _arrow(ax, 40, 40, 32, 40, "results / PDF")

    fig.tight_layout()
    fig.savefig(IMG_DIR / "chapter3_deployment.png", dpi=220)
    plt.close(fig)


def add_heading(document: Document, text: str, level: int = 1):
    document.add_heading(text, level=level)


def add_paragraphs(document: Document, paragraphs):
    for p in paragraphs:
        document.add_paragraph(p)


def build_report_docx():
    doc = Document()

    doc.add_heading("Autism Pre-Screening Tool", 0)
    doc.add_paragraph("Project Report")
    doc.add_paragraph("Faculty of Computer Science and Engineering")
    doc.add_paragraph("Date: April 2026")

    add_heading(doc, "Abstract", 1)
    add_paragraphs(
        doc,
        [
            "This report presents an AI-assisted autism pre-screening system for early risk estimation in toddlers. "
            "The solution combines a photo-based face screening pipeline with Q-CHAT/M-CHAT questionnaire analysis to provide "
            "a practical, parent-facing risk assessment workflow.",
            "A FastAPI backend orchestrates image screening, questionnaire inference, risk explanation generation using a large language model, "
            "and PDF report creation. The system is designed as a screening aid rather than a diagnostic instrument and includes parent-friendly outputs.",
        ],
    )

    add_heading(doc, "Chapter 1: Introduction", 1)
    add_paragraphs(
        doc,
        [
            "Autism Spectrum Disorder (ASD) screening can benefit from early, accessible, and interpretable digital tools. "
            "The project addresses this by implementing a software platform that automates pre-screening workflows while keeping the final decision with healthcare professionals.",
            "Objectives include: (1) collect child information and responses through a simple interface, (2) estimate risk using calibrated machine-learning models, "
            "(3) produce understandable narrative guidance for caregivers, and (4) export documentation for consultation support.",
        ],
    )

    add_heading(doc, "Chapter 2: Literature and Background", 1)
    add_paragraphs(
        doc,
        [
            "The design draws from three practical areas: questionnaire-based autism screening, computer-vision screening pipelines, and explainable AI-assisted reporting.",
            "Questionnaire signals are mapped to model features and combined with demographic factors. The project also applies calibration and threshold tuning "
            "to improve screening sensitivity in high-risk scenarios.",
        ],
    )

    add_heading(doc, "Chapter 3: System Design", 1)
    add_paragraphs(
        doc,
        [
            "Chapter 3 defines architecture, component boundaries, data flow, and deployment assumptions. "
            "The system follows a modular design in which FastAPI serves as the orchestration layer and `src` modules implement domain logic.",
            "Primary modules include face screening (`src/face_screening.py`), questionnaire inference (`src/inference.py`), "
            "LLM report generation (`src/llm_report_groq.py`), and report export (`src/pdf_generator.py`).",
        ],
    )

    figs = [
        ("Figure 1: Overall Architecture", IMG_DIR / "chapter3_architecture.png"),
        ("Figure 2: Use Case View", IMG_DIR / "chapter3_usecase.png"),
        ("Figure 3: Sequence Diagram", IMG_DIR / "chapter3_sequence.png"),
        ("Figure 4: Activity Diagram", IMG_DIR / "chapter3_activity.png"),
        ("Figure 5: Component Diagram", IMG_DIR / "chapter3_component.png"),
        ("Figure 6: Deployment Diagram", IMG_DIR / "chapter3_deployment.png"),
    ]
    for caption, img in figs:
        doc.add_paragraph(caption)
        doc.add_picture(str(img), width=Inches(6.5))
        doc.add_paragraph()

    add_heading(doc, "3.1 Model Description", 2)
    add_paragraphs(
        doc,
        [
            "The questionnaire model uses mapped Q-CHAT/M-CHAT answers and demographics to estimate class probabilities across risk levels "
            "(No Risk, Mild Risk, Moderate Risk, Severe Risk). A per-class threshold policy supports conservative screening behavior.",
            "The face-screening module uses a calibrated CNN pipeline when artifacts are available and can fall back to a legacy EfficientNet checkpoint.",
        ],
    )

    add_heading(doc, "3.2 Functional Requirements", 2)
    add_paragraphs(
        doc,
        [
            "FR-1: System shall accept photo input and run binary face screening.",
            "FR-2: System shall collect Q-CHAT and M-CHAT responses and return risk prediction.",
            "FR-3: System shall generate parent-friendly narrative risk feedback.",
            "FR-4: System shall generate downloadable PDF report.",
            "FR-5: System shall expose API health endpoint.",
        ],
    )

    add_heading(doc, "3.3 Non-Functional Requirements", 2)
    add_paragraphs(
        doc,
        [
            "NFR-1: Clear and simple user interaction for non-technical caregivers.",
            "NFR-2: Robust input validation and stable API behavior.",
            "NFR-3: Maintainability through modular backend structure.",
            "NFR-4: Reproducibility through explicit model artifact management.",
        ],
    )

    add_heading(doc, "Chapter 4: Proposed Solution and Implementation", 1)
    add_paragraphs(
        doc,
        [
            "The implemented solution provides API endpoints for image screening, questionnaire-based prediction, language-adaptive report generation, and PDF export.",
            "Frontend clients communicate with FastAPI endpoints, and backend services coordinate inference and reporting in a deterministic processing pipeline.",
        ],
    )

    add_heading(doc, "Chapter 5: Results and Discussion", 1)
    add_paragraphs(
        doc,
        [
            "The current implementation is functionally complete for screening workflows and report generation. "
            "Script-based tests are available for major modules, though formal CI and automated coverage can be expanded.",
            "Practical strengths include modularity, multilingual report capability, and support for both static-web and Streamlit interfaces.",
        ],
    )

    add_heading(doc, "Chapter 6: Conclusion and Future Work", 1)
    add_paragraphs(
        doc,
        [
            "This project demonstrates a complete autism pre-screening software pipeline that combines computer vision, questionnaire analytics, and narrative reporting.",
            "Future work includes stronger backend authentication, unified frontend standardization, automated test expansion, and longitudinal evaluation with clinically curated datasets.",
        ],
    )

    add_heading(doc, "References", 1)
    add_paragraphs(
        doc,
        [
            "FastAPI Documentation. https://fastapi.tiangolo.com/",
            "scikit-learn Documentation. https://scikit-learn.org/",
            "PyTorch Documentation. https://pytorch.org/docs/",
            "Q-CHAT and M-CHAT screening literature (clinical references to be finalized with supervisor).",
        ],
    )

    doc.save(DOCX_PATH)


def main():
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    create_architecture()
    create_usecase()
    create_sequence()
    create_activity()
    create_component()
    create_deployment()
    build_report_docx()
    print(f"Generated report: {DOCX_PATH}")
    print(f"Generated images in: {IMG_DIR}")


if __name__ == "__main__":
    main()
